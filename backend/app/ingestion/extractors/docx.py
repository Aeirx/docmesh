"""DOCX extraction via python-docx.

Structure comes from paragraph styles: "Heading N" styles drive the section path,
"Title" becomes the document title. Everything else is a body paragraph.

Documented limitations (honest Phase 2 scope cuts): ``page_number``/``page_count``
are None — DOCX has no fixed pagination, pages are a renderer artifact; tables and
text boxes are skipped (``doc.paragraphs`` only).

python-docx is imported lazily inside extract() — same rationale as pdf.py.
"""

from pathlib import Path

from app.ingestion.extractors.base import (
    ExtractedDoc,
    ExtractionError,
    _DocBuilder,
    _SectionTracker,
)


def extract(path: Path) -> ExtractedDoc:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"Cannot open DOCX: {exc}") from exc

    builder = _DocBuilder()
    tracker = _SectionTracker()
    title: str | None = None

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style is not None else None) or ""

        if style.startswith("Heading"):
            last = style.rsplit(" ", 1)[-1]
            level = min(int(last), 3) if last.isdigit() else 1
            tracker.on_heading(level, text)
            builder.add(
                text, page=None, section=tracker.path(), is_heading=True, heading_level=level
            )
        elif style == "Title" and title is None:
            title = text
            tracker.on_heading(1, text)
            builder.add(text, page=None, section=tracker.path(), is_heading=True, heading_level=1)
        else:
            builder.add(text, page=None, section=tracker.path())

    return builder.build(page_count=None, title=title)
